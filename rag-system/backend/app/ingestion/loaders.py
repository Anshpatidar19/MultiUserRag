"""
loaders.py

One function per source format, each returning plain extracted text (or
raising a `LoaderError` with a specific, user-facing reason). Deliberately
NOT doing chunking or embedding here -- a loader's only job is "get clean
text out of this file," so it can be tested and reasoned about in
isolation from the rest of the pipeline.

Web page scraping is intentionally NOT implemented here (out of scope
per spec) -- there is no `load_url` function, and `pipeline.py` never
routes a "paste a URL" request into this module for arbitrary web pages.
YouTube is the one URL-based exception, since it's transcript extraction
via the YouTube API, not generic page scraping.
"""

import io
import csv
import pdfplumber
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document as DocxDocument
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound

class LoaderError(Exception):
    """Raised with a message safe to show directly to the user."""


_OCR_CONFIG = "--psm 6"


def load_pdf(file_bytes: bytes, describe_page_fn=None) -> str:
   
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) == 0:
                raise LoaderError("PDF has no pages.")

            text_parts: list[str] = []
            ocr_needed_pages: list[int] = []
            for i, page in enumerate(pdf.pages):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    text_parts.append(page_text)
                else:
                    text_parts.append("")  # placeholder, filled in by OCR pass below
                    ocr_needed_pages.append(i)
    except LoaderError:
        raise
    except Exception as exc:  # noqa: BLE001 - surface as a user-facing loader error
        raise LoaderError(f"Could not read this PDF: {exc}") from exc

    if ocr_needed_pages:
        images = convert_from_bytes(file_bytes)
        for i in ocr_needed_pages:
            if i >= len(images):
                continue
            page_image = images[i]
            ocr_text = pytesseract.image_to_string(page_image, config=_OCR_CONFIG).strip()

            if describe_page_fn is not None:
                buf = io.BytesIO()
                page_image.save(buf, format="PNG")
                try:
                    vision_text = describe_page_fn(buf.getvalue(), ocr_text=ocr_text)
                except Exception:  # noqa: BLE001 - never lose the raw OCR text over a vision-call failure
                    vision_text = ""
                text_parts[i] = "\n\n".join(t for t in (ocr_text, vision_text) if t.strip())
            else:
                text_parts[i] = ocr_text

    full_text = "\n\n".join(t for t in text_parts if t)
    if not full_text.strip():
        raise LoaderError(
            "Could not extract any text from this PDF, even with OCR. "
            "The file may be corrupted or blank."
        )
    return full_text


def load_image(file_bytes: bytes) -> str:
    """
    OCR first (cheap, catches printed/handwritten text in the photo).
    The caller's pipeline always follows this up with a vision-model
    description (see pipeline.py) -- OCR text alone can't count people
    in a photo or reliably read a table/marksheet's row/column
    structure, so it's treated as a helpful hint for the vision model
    rather than the final answer.
    """
    image = Image.open(io.BytesIO(file_bytes))
    ocr_text = pytesseract.image_to_string(image, config=_OCR_CONFIG).strip()
    return ocr_text  # may be empty string; pipeline decides what to do next


def load_csv(file_bytes: bytes) -> str:
    """
    Render rows as "col: value, col: value" lines rather than raw CSV,
    so downstream chunking/embedding operates on natural-language-ish
    text instead of comma-separated tokens that embed poorly.
    """
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    if reader.fieldnames is None:
        raise LoaderError("CSV has no header row or is empty.")

    rows_text = []
    for row in reader:
        line = ", ".join(f"{k}: {v}" for k, v in row.items() if v not in (None, ""))
        if line:
            rows_text.append(line)

    if not rows_text:
        raise LoaderError("CSV parsed but contained no non-empty rows.")
    return "\n".join(rows_text)


def load_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables carry real content in DOCX reports -- don't silently drop them.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise LoaderError("DOCX file contained no extractable text.")
    return "\n\n".join(paragraphs)


def _extract_video_id(url: str) -> str:
    import re

    patterns = [
        r"(?:v=|/)([0-9A-Za-z_-]{11}).*",
        r"youtu\.be/([0-9A-Za-z_-]{11})",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise LoaderError("Could not parse a video ID from that YouTube URL.")


def load_youtube(url: str) -> str:
    video_id = _extract_video_id(url)
    try:
        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)
    except TranscriptsDisabled as exc:
        raise LoaderError("Transcripts are disabled for this video.") from exc
    except NoTranscriptFound as exc:
        raise LoaderError("No transcript is available for this video.") from exc
    except Exception as exc:  # noqa: BLE001
        # YouTube has been increasingly rate-limiting/blocking automated
        # transcript requests (a known, ongoing issue -- see
        # https://github.com/jdepoix/youtube-transcript-api/issues/429).
        # Surface this as a clear, actionable message instead of a raw
        # XML parse error, since it's not something the user's URL did wrong.
        raise LoaderError(
            "Could not fetch this video's transcript -- YouTube may be "
            "rate-limiting automated requests right now. Try again in a "
            "few minutes, or try a different video."
        ) from exc

    full_text = " ".join(snippet.text for snippet in transcript if snippet.text.strip())
    if not full_text.strip():
        raise LoaderError("Transcript was empty.")
    return full_text